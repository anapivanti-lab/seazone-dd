"""PARECER JURÍDICO único da Due Diligence (formato do setor de Franquias).

Cada papel (Franquia / Representante legal / Operador) é processado num "run" e,
ao gerar o parecer, salva um snapshot (_dd_<papel>.json) na pasta da DD. O parecer
lê TODOS os snapshots da pasta e monta UM ÚNICO documento Word (.docx) editável com
uma seção por papel + conclusão combinada (o risco mais grave entre todos).

Conclusão em 4 níveis (sobre o conjunto): nada / achado leve / risco financeiro real
/ risco irreversível (criminal).
"""
from __future__ import annotations

import json
import re
import urllib.parse
from datetime import datetime
from pathlib import Path

from . import drive_links, ia_local
from .extrator import _texto_documento
from .pep import checar as checar_pep
from .storage import _slug

LIMITE_DEBITO = 10000.0

TEXTO_OK = ("Após realizada a análise da documentação elencada acima, certificou-se que, "
            "estritamente em relação aos documentos apresentados pela franquia e os emitidos pelo "
            "setor Jurídico da Seazone, estes não apresentam riscos, impedimentos ou restrições que "
            "possam influenciar na concessão da franquia. Logo, não há impedimentos, podendo seguir "
            "com a contratação.")

_MESES = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho",
          "agosto", "setembro", "outubro", "novembro", "dezembro"]
_ORDEM = {"Franquia": 0, "Representante legal": 1, "Representante legal e Operador": 1, "Operador": 2}
_CORES = {"ALTO": "#c0392b", "MÉDIO": "#b8860b", "BAIXO": "#1a7d3c"}


# ---------------------------------------------------------------- utilidades
def _fmt(doc: str, pj: bool) -> str:
    d = re.sub(r"\D", "", doc or "")
    if pj and len(d) == 14:
        return f"{d[:2]}.{d[2:5]}.{d[5:8]}/{d[8:12]}-{d[12:]}"
    if not pj and len(d) == 11:
        return f"{d[:3]}.{d[3:6]}.{d[6:9]}-{d[9:]}"
    return doc or "—"


def _reais(v: float) -> str:
    return ("R$ " + f"{v:,.2f}").replace(",", "X").replace(".", ",").replace("X", ".")


def _genero(nome: str) -> str:
    p = (nome or "").strip().split()
    return "f" if p and p[0].lower().endswith("a") else "m"


def _data_extenso() -> str:
    d = datetime.now()
    return f"Florianópolis/SC, {d.day} de {_MESES[d.month - 1]} de {d.year}."


def _classificar(caminho: str) -> str:
    t = _texto_documento(caminho).lower()
    if not t.strip():
        return "indeterminado"
    if any(k in t for k in ("insuficientes para emitir", "não foi possível emitir",
                            "nao foi possivel emitir", "não é possível emitir", "nao e possivel emitir")):
        return "positiva"  # comunicado da Receita = CND positiva
    if "positiva com efeito" in t or "positiva com efeitos" in t:
        return "negativa"
    if any(k in t for k in ("nada consta", "não constam", "nao constam", "negativ", "inexist")):
        return "negativa"
    if any(k in t for k in ("positiv", "constam débitos", "constam debitos", "consta o débito", "em aberto")):
        return "positiva"
    return "indeterminado"


def _valor_protestos(caminho: str) -> float:
    vals = []
    for m in re.findall(r"R\$\s*([\d.]+,\d{2})", _texto_documento(caminho)):
        try:
            vals.append(float(m.replace(".", "").replace(",", ".")))
        except Exception:
            pass
    return sum(vals)


def _fatos_curto(fatos: str, limite: int = 820) -> str:
    """Resumo objetivo dos fatos para o parecer: remove o 'Consta...' inicial, o
    juridiquês mais pesado e citações/coordenadas, e limita a poucas frases-chave."""
    if not fatos:
        return ""
    t = re.sub(r"^\s*Consta d[oa][^,]{0,45},\s*", "", fatos, flags=re.I)
    t = re.sub(r"^\s*Consta que,?\s*", "", t, flags=re.I)
    for a, b in [
        (r"em comunh[ãa]o de esfor[çc]os e unidade de des[íi]gnios com", "em conjunto com"),
        (r",?\s*aderindo [àa] conduta inicial daquele,?", ""),
        (r"\s*e?\s*agindo com animus rem sibi habendi\s*\([^)]*\)", ""),
        (r"inverteram o t[íi]tulo da posse", "ficaram com o bem indevidamente"),
        (r"promove[a-z]+ a desconex[ãa]o do", "desconectaram o"),
        (r"deixando de restituir o autom[óo]vel", "não devolvendo o veículo"),
        (r"\(localizado nas coordenadas[^)]*\)?", ""),
        (r"\((?:evento|fls?\.|p\.|INQ|n\.|seq|art)[^)]*\)?", ""),
        (r",?\s*de placas? [A-Z0-9-]+", ""),
        (r"\bpel[oa]s? denunciad[oa]s?\b", "por"),
        (r"\bos? denunciados?\b", ""),
        (r"\ba denunciada\b", ""),
        (r"consumando-se a apropriação do bem alheio móvel de que detinham a posse",
         "consumando-se a apropriação do veículo"),
        (r"\s*\(\s*$", ""),
        (r"habilita[çc][ãa]o\d", "habilitação"),
    ]:
        t = re.sub(a, b, t, flags=re.I)
    t = re.sub(r"\s+", " ", t).strip()
    frases = re.split(r"(?<=[.;])\s+", t)
    # mantém a 1ª frase (contexto) + as frases-chave (o que de fato aconteceu),
    # descartando enrolação — assim o resumo é curto mas pega o essencial.
    CHAVE = ("não dev", "nao dev", "restitu", "apropria", "rastreador", "desconect", "recuper",
             "polícia", "policia", "flagrante", "devolv", "posse indevida", "subtra", "furt", "estelionat")
    sel = [fr for i, fr in enumerate(frases) if i == 0 or any(k in fr.lower() for k in CHAVE)]
    out = " ".join(sel or frases[:3])
    out = re.sub(r"\s+([,.;])", r"\1", re.sub(r"\s{2,}", " ", out)).strip(" ,;")
    if out:
        out = out[0].upper() + out[1:]
        if not out.endswith((".", "…")):
            out += "."
    return out


def _qual_franquia(ctx, dados) -> str:
    sede = ctx.endereco or dados.get("endereco") or "endereço não informado"
    return (f"{ctx.nome or 'Razão social não informada'}, pessoa jurídica de direito privado, "
            f"inscrita no CNPJ sob o nº {_fmt(ctx.documento, True)}, com sede em {sede}.")


def _qual_operador(ctx, dados) -> str:
    g = _genero(ctx.nome)
    bras = "brasileira" if g == "f" else "brasileiro"
    insc = "inscrita" if g == "f" else "inscrito"
    domic = "domiciliada" if g == "f" else "domiciliado"
    ec = (ctx.estado_civil or "").strip()
    if ec.endswith("(a)"):
        base = ec[:-3]
        ec = (base[:-1] + "a") if g == "f" else base
    ec = ec.lower() if ec else "estado civil não informado"
    end = ctx.endereco or "endereço não informado"
    return (f"{ctx.nome or 'Nome não informado'}, {bras}, {ec}, {insc} no CPF sob o nº "
            f"{_fmt(ctx.documento, False)}, residente e {domic} em {end}.")


# ---------------------------------------------------------------- análise por entidade
def _analisar_entidade(job) -> dict:
    ctx = job.ctx
    dados = getattr(job, "cnpj_dados", {}) or {}
    pj = ctx.tipo.value == "PJ"
    papel = ctx.papel or ("Franquia" if pj else "Operador")
    de_quem = "da Franquia" if pj else ("da " + papel if papel.lower().startswith("represent") and _genero(ctx.nome) == "f"
                                        else "do " + papel)
    sujeito = "a Franquia" if pj else f"o(a) {papel}"

    GRUPOS = ("Federais", "Estaduais", "Justiça Estadual", "Justiça Federal", "Municipais")
    certs = [{"nome": p.nome, "classe": _classificar(p.arquivo)}
             for p in job.passos if p.arquivo and p.grupo in GRUPOS and "Protesto" not in p.nome]
    positivas = [c["nome"] for c in certs if c["classe"] == "positiva"]
    indet = [c["nome"] for c in certs if c["classe"] == "indeterminado"]

    passo_prot = next((p for p in job.passos if "Protesto" in p.nome), None)
    prot_classe = _classificar(passo_prot.arquivo) if (passo_prot and passo_prot.arquivo) else "sem"
    prot_positiva = prot_classe == "positiva"
    total_prot = _valor_protestos(passo_prot.arquivo) if (prot_positiva and passo_prot.arquivo) else 0.0

    procs = job.processos or []
    crim = any(pr.get("criminal") for pr in procs)
    total_proc = sum(pr.get("valor_maximo") or 0 for pr in procs)

    situacao = dados.get("situacao", "")
    irregular = situacao if (situacao and situacao != "ATIVA") else ""
    pendencias = [p.nome for p in job.passos if p.grupo == "Você fornece" and not p.arquivo]

    pep = checar_pep([], cpf=ctx.documento) if not pj else checar_pep(dados.get("socios", []))
    pep_nomes = [f"{p['nome']} ({p['funcao']})" for p in (pep.get("pep") or [])]

    if not certs:
        t_cert = "Não há certidões coletadas até o momento."
    elif positivas:
        t_cert = "Constam certidões POSITIVAS em nome " + de_quem + ": " + "; ".join(positivas) + "."
    elif indet:
        t_cert = ("As certidões foram coletadas. Confira manualmente as que não puderam ser lidas "
                  "automaticamente: " + "; ".join(indet) + ".")
    else:
        t_cert = "Todas as certidões estão regulares (negativas ou positivas com efeito de negativa)."

    if not procs:
        t_proc = f"Não foram localizados registros de ações judiciais ativas em nome {de_quem}."
    else:
        itens = []
        for pr in procs:
            resumo_ia = ia_local.resumir_processo(pr, sujeito)  # IA local (se disponível)
            if resumo_ia:
                itens.append(resumo_ia)
                continue
            num = pr.get("numero") or "s/ número"
            classe = (pr.get("classe") or "").strip()
            tipo = "criminal" if pr.get("criminal") else (classe.lower() if classe else "judicial")
            pd = (pr.get("papel_dd") or "").lower()
            if any(k in pd for k in ("passivo", "réu", "reu", "requerid", "executad", "denunciad")):
                polo = "no polo passivo"
            elif any(k in pd for k in ("ativo", "autor", "exequente", "reclamante")):
                polo = "no polo ativo"
            else:
                polo = "no processo"
            frase = f"Foi localizado o processo {tipo} de nº {num}, em que {sujeito} figura {polo}"
            if pr.get("assunto"):
                frase += f", relativo a {pr['assunto'].split(',')[0].strip().lower()}"
            frase += ". "
            if pr.get("valor_maximo"):
                frase += f"Valor/débito de {_reais(pr['valor_maximo'])}. "
            if pr.get("sentenca"):
                frase += f"Desfecho: {pr['sentenca'].get('resultado')}. "
            fatos = _fatos_curto(pr.get("fatos") or "")
            if fatos:
                frase += fatos
            itens.append(frase.strip())
        t_proc = " ".join(itens)

    if prot_classe == "negativa":
        t_prot = f"Não há registros de títulos protestados em nome {de_quem}."
    elif prot_positiva:
        v = f", que totalizam {_reais(total_prot)}" if total_prot else ""
        t_prot = f"Constam títulos protestados em nome {de_quem}{v}."
    else:
        t_prot = "Certidão de protestos pendente de emissão/verificação."

    t_pend = ("Restam pendentes de envio: " + "; ".join(pendencias) + ".") if pendencias else ""

    qual = _qual_franquia(ctx, dados) if pj else _qual_operador(ctx, dados)
    return {
        "papel": papel, "ordem": _ORDEM.get(papel, 3), "tipo": ctx.tipo.value,
        "nome": ctx.nome, "qualificacao": qual,
        "certidoes_txt": t_cert, "processos_txt": t_proc, "protestos_txt": t_prot, "pendencias_txt": t_pend,
        "documento": ctx.documento,
        "cnae": dados.get("cnae_codigo", ""), "cnae_desc": dados.get("cnae_descricao", ""),
        "itens": [{"nome": p.nome, "grupo": p.grupo, "status": p.status,
                   "arquivo": (Path(p.arquivo).name if p.arquivo else "")} for p in job.passos],
        "positivas": positivas, "prot_positiva": prot_positiva, "crim": crim, "procs": bool(procs),
        "situacao_irregular": irregular, "pendencias": pendencias,
        "total_proc": total_proc, "total_prot": total_prot, "pep": pep_nomes,
    }


def _snapshot(pasta: Path, ent: dict) -> None:
    try:
        (pasta / f"_dd_{_slug(ent['papel'])}.json").write_text(
            json.dumps(ent, ensure_ascii=False), encoding="utf-8")
    except Exception:
        pass


def _carregar_entidades(pasta: Path) -> list:
    ents = []
    for f in pasta.glob("_dd_*.json"):
        try:
            ents.append(json.loads(f.read_text(encoding="utf-8")))
        except Exception:
            pass
    ents.sort(key=lambda e: e.get("ordem", 3))
    return ents


# ---------------------------------------------------------------- conclusão combinada
def _recs_entidade(e: dict) -> list:
    lab = e["papel"]
    out = []
    if e["positivas"]:
        out.append(f"Regularizar a situação fiscal e reverter/esclarecer a(s) certidão(ões) positiva(s) ({lab})")
    if e["procs"]:
        v = f" (débito/valor não atualizado de {_reais(e['total_proc'])})" if e.get("total_proc") else ""
        out.append(f"Regularizar a situação decorrente das ações judiciais{v} ({lab})")
    if e["prot_positiva"]:
        v = f" ({_reais(e['total_prot'])})" if e.get("total_prot") else ""
        out.append(f"Pagamento/baixa dos títulos protestados{v} ({lab})")
    if e["situacao_irregular"]:
        out.append(f"Regularização da situação cadastral ({e['situacao_irregular']}) do CNPJ ({lab})")
    if e.get("pep"):
        out.append(f"Verificar a exposição política identificada ({lab}): " + "; ".join(e["pep"]))
    if e["pendencias"]:
        out.append(f"Envio dos documentos pendentes ({lab}): " + "; ".join(e["pendencias"]))
    return out


def _conclusao(entidades: list) -> dict:
    crim = any(e["crim"] for e in entidades)
    positivas = any(e["positivas"] for e in entidades)
    prot = any(e["prot_positiva"] for e in entidades)
    procs = any(e["procs"] for e in entidades)
    irregular = any(e["situacao_irregular"] for e in entidades)
    achado = any(e["positivas"] or e["prot_positiva"] or e["procs"] or e["situacao_irregular"]
                 or e["pendencias"] or e.get("pep") for e in entidades)
    total = sum((e.get("total_proc") or 0) + (e.get("total_prot") or 0) for e in entidades)
    n_pos = sum(1 for e in entidades if e["positivas"])
    grave = total >= LIMITE_DEBITO or n_pos >= 2 or (prot and positivas) or (procs and total >= LIMITE_DEBITO)
    recs = [r for e in entidades for r in _recs_entidade(e)]

    if crim:
        return {"risco": "ALTO", "categoria": "Processual",
                "texto": ("A identificação de processo(s) criminal(is) representa fator de atenção no contexto "
                          "da análise de risco, apresentando riscos relevantes, principalmente reputacionais."),
                "recs": ["Não aprovar a concessão da franquia."], "fecho": ""}
    if grave:
        tipos = []
        if prot:
            tipos.append("protestos")
        if positivas:
            tipos.append("certidões positivas")
        if procs:
            tipos.append("execuções")
        if irregular:
            tipos.append("situação cadastral irregular")
        resumo = (", ".join(tipos[:-1]) + " e " + tipos[-1]) if len(tipos) > 1 else (tipos[0] if tipos else "os achados")
        return {"risco": "ALTO", "categoria": "Financeiro",
                "texto": resumo[0].upper() + resumo[1:] + " indicam fragilidade de liquidez e histórico de inadimplência relevante.",
                "recs": [r if r.endswith(".") else r + "." for r in recs],
                "fecho": "Com a conclusão das recomendações acima, podemos seguir com a contratação."}
    if achado:
        return {"risco": "MÉDIO", "categoria": "",
                "texto": TEXTO_OK + " Porém, recomendamos: " + "; ".join(recs) + ".",
                "recs": [], "fecho": ""}
    return {"risco": "BAIXO", "categoria": "", "texto": TEXTO_OK, "recs": [], "fecho": ""}


ARQ_PARECER = "Parecer_Juridico_DD.docx"
ARQ_RELATORIO = "Relatorio_DD.pdf"


def _criterios(entidades: list) -> list:
    """Os 6 critérios de risco -> (texto, SIM/NÃO/None, observação)."""
    pos = any(e["positivas"] for e in entidades)
    crim = any(e["crim"] for e in entidades)
    prot = any(e["prot_positiva"] for e in entidades)
    pep = any(e.get("pep") for e in entidades)
    civel = any(e["procs"] and not e["crim"] for e in entidades)
    cnae = next((f'{e.get("cnae", "")} {e.get("cnae_desc", "")}'.strip() for e in entidades if e.get("cnae")), "")
    return [
        ("Todas as certidões são negativas (ou positivas com efeito de negativa)?", not pos, ""),
        ("Inexiste processo criminal envolvendo o representante/franquia?", not crim, ""),
        ("O CNAE utilizado pela franquia está correto?", None, cnae or "conferir manualmente"),
        ("O representante legal não é pessoa politicamente exposta (PEP)?", not pep, ""),
        ("Inexistem protestos em nome da franquia/representante?", not prot, ""),
        ("Inexistem processos cíveis relacionados ao objeto da franquia?", not civel, ""),
    ]


# ---------------------------------------------------------------- montar + salvar
def gerar(job) -> dict:
    ctx = job.ctx
    pasta = ctx.pasta_saida
    ent = _analisar_entidade(job)
    _snapshot(pasta, ent)                 # salva o snapshot deste papel
    entidades = _carregar_entidades(pasta) or [ent]
    concl = _conclusao(entidades)

    titulo = (next((e.get("nome") for e in entidades if "operador" in e["papel"].lower() and e.get("nome")), None)
              or (entidades[0].get("nome") if entidades else None) or ctx.nome or ctx.documento)
    d = {
        "titulo": titulo, "id_suporte": ctx.id_suporte, "entidades": entidades,
        "risco": concl["risco"], "concl_categoria": concl["categoria"], "concl_texto": concl["texto"],
        "recomendacoes": concl["recs"], "concl_fecho": concl["fecho"],
        "gerado_em": datetime.now().strftime("%d/%m/%Y %H:%M"), "data_extenso": _data_extenso(),
    }
    # Links clicáveis do Drive para as certidões/arquivos já sincronizados
    nomes = [it["arquivo"] for e in entidades for it in e.get("itens", []) if it["arquivo"]]
    try:
        d["urls"], d["url_pasta"] = drive_links.urls_da_pasta(pasta.name, nomes)
    except Exception:
        d["urls"], d["url_pasta"] = {}, ""

    # HTML do relatório (vira PDF no endpoint, via Chromium — formato compartilhável)
    d["relatorio_html"] = _relatorio_html(d, _criterios(entidades))
    d["html"] = _pagina_html(d)
    d["arquivo"] = str(_salvar_docx(pasta, d))
    return d


async def render_relatorio_pdf(pasta: Path, html: str) -> None:
    """Renderiza o relatório (HTML) em PDF (Relatorio_DD.pdf) — formato fácil de
    compartilhar por link (abre no navegador pra qualquer pessoa)."""
    if not html:
        return
    from playwright.async_api import async_playwright
    try:
        pw = await async_playwright().start()
        try:
            nav = await pw.chromium.launch()
            page = await nav.new_page()
            await page.set_content(html, wait_until="load")
            await page.pdf(path=str(pasta / ARQ_RELATORIO), format="A4", print_background=True,
                           margin={"top": "1.2cm", "bottom": "1.2cm", "left": "1.2cm", "right": "1.2cm"})
            await nav.close()
        finally:
            await pw.stop()
    except Exception:
        try:  # se o PDF falhar, salva ao menos o HTML
            (pasta / "Relatorio_DD.html").write_text(html, encoding="utf-8")
        except Exception:
            pass


# ---------------------------------------------------------------- WORD (.docx)
def _rgb(hexcor):
    from docx.shared import RGBColor
    h = hexcor.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _hyperlink(par, url, texto, cor="0B4F6C"):
    from docx.oxml.shared import OxmlElement, qn
    rid = par.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    h = OxmlElement("w:hyperlink")
    h.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), cor)
    rpr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    run.append(rpr)
    tt = OxmlElement("w:t")
    tt.text = texto
    run.append(tt)
    h.append(run)
    par._p.append(h)


def _salvar_docx(pasta: Path, d) -> Path:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    AZUL = _rgb("#0b4f6c")
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    for m in doc.sections:
        m.left_margin = m.right_margin = Pt(64)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cab = "PARECER JURÍDICO — DUE DILIGENCE" + (f"  (#{d['id_suporte']})" if d["id_suporte"] else "")
    rt = t.add_run(cab)
    rt.bold = True
    rt.font.size = Pt(15)
    rt.font.color.rgb = AZUL
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run(d["titulo"])
    rs.bold = True
    rs.font.size = Pt(12)
    rk = doc.add_paragraph()
    rk.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = rk.add_run(f"Risco: {d['risco']}")
    rr.bold = True
    rr.font.color.rgb = _rgb(_CORES.get(d["risco"], "#0b4f6c"))

    def secao(txt):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        r = p.add_run(txt)
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = AZUL

    def label(lbl, texto, recuo=False):
        p = doc.add_paragraph()
        if recuo:
            p.paragraph_format.left_indent = Pt(16)
        p.add_run(lbl + " ").bold = True
        if texto:
            p.add_run(texto)

    secao("1. Qualificação")
    for e in d["entidades"]:
        label(e["papel"] + ":", e["qualificacao"])

    secao("2. Documentos analisados")
    p = doc.add_paragraph()
    if d.get("url_pasta"):
        p.add_run(f"As certidões e o relatório completo (“{ARQ_RELATORIO}”) estão na pasta da DD no Drive: ")
        _hyperlink(p, d["url_pasta"], "abrir a pasta")
        p.add_run(".")
    else:
        p.add_run(f"As certidões e documentos analisados constam no Relatório da DD "
                  f"(arquivo “{ARQ_RELATORIO}”), salvo na mesma pasta desta Due Diligence.")

    secao("3. Parecer")
    for e in d["entidades"]:
        pe = doc.add_paragraph()
        re_ = pe.add_run(e["papel"])
        re_.bold = True
        re_.underline = True
        label("Certidões:", e["certidoes_txt"], recuo=True)
        label("Processos:", e["processos_txt"], recuo=True)
        label("Protestos:", e["protestos_txt"], recuo=True)
        if e["pendencias_txt"]:
            label("Pendências:", e["pendencias_txt"], recuo=True)

    secao("4. Conclusão")
    if d["concl_categoria"]:
        label(d["concl_categoria"] + ":", d["concl_texto"])
        doc.add_paragraph().add_run("Recomendação:").bold = True
        for r in d["recomendacoes"]:
            doc.add_paragraph(r, style="List Number")
        if d["concl_fecho"]:
            doc.add_paragraph(d["concl_fecho"])
    else:
        doc.add_paragraph(d["concl_texto"])

    doc.add_paragraph()
    doc.add_paragraph(d["data_extenso"])
    doc.add_paragraph(f"Relatório completo da DD (checklist e certidões): arquivo “{ARQ_RELATORIO}”, na mesma pasta.")

    destino = pasta / ARQ_PARECER
    doc.save(str(destino))
    return destino


# ---------------------------------------------------------------- HTML (prévia)
_ESTILO = """
body{font-family:Georgia,'Times New Roman',serif;color:#1a2332;max-width:820px;margin:1rem auto;line-height:1.65;padding:0 1.4rem;background:#fff}
.cab{text-align:center;border-bottom:3px solid #0b4f6c;padding-bottom:.8rem;margin-bottom:1rem}
.cab h1{color:#0b4f6c;font-size:1.25rem;margin:.2rem 0}
.titulo{font-weight:600;font-size:1.05rem;margin:.2rem 0}
h1{font-size:1.15rem;margin:.5rem 0 .8rem;font-weight:700}
h2{font-size:1.0rem;color:#0b4f6c;border-bottom:1px solid #d3dde3;padding-bottom:.2rem;margin:1.3rem 0 .5rem}
p{text-align:justify;margin:.35rem 0}.q b{color:#0b4f6c}
a{color:#0b6;font-weight:600}
.ent{font-weight:700;margin-top:.7rem;text-decoration:underline}
.sub{font-weight:700;margin:.45rem 0 .1rem;padding-left:.6rem}
.badge{display:inline-block;padding:.25rem 1rem;border-radius:999px;color:#fff;font-weight:700;font-size:.85rem;margin-top:.3rem}
ul,ol{margin:.3rem 0 .3rem 1.2rem}.rec li{margin:.25rem 0}
.fim{margin-top:1.2rem;color:#333}
"""


def _corpo_html(d) -> str:
    cor = _CORES.get(d["risco"], "#1a7d3c")
    quals = "".join(f'<p class="q"><b>{e["papel"]}:</b> {e["qualificacao"]}</p>' for e in d["entidades"])
    blocos = ""
    for e in d["entidades"]:
        pend = f'<p class="sub">Pendências:</p><p>{e["pendencias_txt"]}</p>' if e["pendencias_txt"] else ""
        blocos += (f'<p class="ent">{e["papel"]}</p>'
                   f'<p class="sub">Certidões:</p><p>{e["certidoes_txt"]}</p>'
                   f'<p class="sub">Processos:</p><p>{e["processos_txt"]}</p>'
                   f'<p class="sub">Protestos:</p><p>{e["protestos_txt"]}</p>{pend}')
    if d["concl_categoria"]:
        recs = "".join(f"<li>{r}</li>" for r in d["recomendacoes"])
        fecho = f"<p>{d['concl_fecho']}</p>" if d["concl_fecho"] else ""
        concl = (f'<p><b>{d["concl_categoria"]}:</b> {d["concl_texto"]}</p>'
                 f'<p class="sub" style="padding-left:0">Recomendação:</p><ol class="rec">{recs}</ol>{fecho}')
    else:
        concl = f"<p>{d['concl_texto']}</p>"
    idtxt = f" · #{d['id_suporte']}" if d["id_suporte"] else ""
    if d.get("url_pasta"):
        docs2 = (f'As certidões e o relatório completo estão na '
                 f'<a href="{d["url_pasta"]}" target="_blank">pasta da DD no Drive ↗</a>.')
    else:
        docs2 = f'As certidões e documentos analisados constam no <b>Relatório da DD</b> (arquivo “{ARQ_RELATORIO}”), na mesma pasta.'
    return f"""
    <div class="cab">
      <h1>Parecer Jurídico — Due Diligence{idtxt}</h1>
      <p class="titulo">{d['titulo']}</p>
      <span class="badge" style="background:{cor}">Risco: {d['risco']}</span>
    </div>
    <h2>1. Qualificação</h2>{quals}
    <h2>2. Documentos analisados</h2>
    <p>{docs2}</p>
    <h2>3. Parecer</h2>{blocos}
    <h2>4. Conclusão</h2>{concl}
    <p class="fim">{d['data_extenso']}</p>
    """


def _pagina_html(d) -> str:
    return (f'<!doctype html><html lang="pt-br"><head><meta charset="utf-8">'
            f'<title>Parecer Jurídico — {d["titulo"]}</title>'
            f"<style>{_ESTILO}</style></head><body>{_corpo_html(d)}</body></html>")


# ---------------------------------------------------------------- RELATÓRIO (HTML)
_SIT = {
    "sucesso": ("✅", "Emitida", "#1a7d3c"), "enviado": ("✅", "Enviada", "#1a7d3c"),
    "manual": ("📤", "Falta enviar", "#b8860b"), "aberta": ("📂", "Em aberto", "#b8860b"),
    "pendente": ("⏳", "Pendente", "#b8860b"), "aguardando": ("•", "Pendente", "#8a97a3"),
    "local": ("📍", "Falta UF/cidade", "#b8860b"), "erro": ("❌", "Erro", "#c0392b"),
}

_RELATORIO_ESTILO = """
body{font-family:'Segoe UI',Arial,sans-serif;color:#1a2332;max-width:900px;margin:1.2rem auto;padding:0 1.4rem;background:#f4f6f8}
.card{background:#fff;border-radius:12px;box-shadow:0 1px 4px rgba(0,0,0,.08);padding:1.4rem 1.6rem;margin-bottom:1.2rem}
.cab{text-align:center;border-bottom:3px solid #0b4f6c;padding-bottom:.9rem}
.cab h1{color:#0b4f6c;margin:.2rem 0;font-size:1.3rem}
.cab .titulo{font-size:1.05rem;font-weight:600;margin:.2rem 0}
.cab .data{color:#8a97a3;font-size:.85rem}
.badge{display:inline-block;padding:.3rem 1rem;border-radius:999px;color:#fff;font-weight:700;font-size:.9rem;margin-top:.4rem}
.ident-item{margin:.25rem 0}.muted{color:#8a97a3}
h2{color:#0b4f6c;font-size:1.05rem;margin:.2rem 0 .6rem;border-bottom:1px solid #e3e6ea;padding-bottom:.25rem}
table.rel{border-collapse:collapse;width:100%}
table.rel th,table.rel td{border:1px solid #e3e6ea;padding:.5rem .6rem;text-align:left;font-size:.9rem}
table.rel th{background:#eef3f7;color:#0b4f6c}
table.rel a{color:#0b6;font-weight:600;text-decoration:none}table.rel a:hover{text-decoration:underline}
.arq{font-family:Consolas,monospace;font-size:.8rem;color:#5a6b7a;word-break:break-all}
.sub{font-weight:700;margin:.5rem 0 .2rem}
ol{margin:.3rem 0 .3rem 1.2rem}ol li{margin:.25rem 0}
.btn-parecer{display:inline-block;background:#0b4f6c;color:#fff;padding:.6rem 1.2rem;border-radius:8px;font-weight:700;text-decoration:none;margin-top:.4rem}
.rodape{color:#8a97a3;font-size:.78rem;text-align:center;margin-top:.5rem}
"""


def _relatorio_html(d, criterios) -> str:
    cor = _CORES.get(d["risco"], "#1a7d3c")
    idtxt = f" · #{d['id_suporte']}" if d["id_suporte"] else ""
    pasta_btn = (f'<a class="btn-parecer" href="{d.get("url_pasta")}" target="_blank">📁 Abrir a pasta da DD no Drive</a>'
                 if d.get("url_pasta") else "")

    ident = ""
    for e in d["entidades"]:
        lbl = "CNPJ" if e["tipo"] == "PJ" else "CPF"
        ident += (f'<div class="ident-item"><b>{e["papel"]}:</b> {e.get("nome") or "—"} '
                  f'<span class="muted">({lbl} {_fmt(e.get("documento", ""), e["tipo"] == "PJ")})</span></div>')

    tabelas = ""
    for e in d["entidades"]:
        linhas = ""
        for it in e.get("itens", []):
            ic, txt, c = _SIT.get(it["status"], ("•", it["status"], "#8a97a3"))
            url = d.get("urls", {}).get(it["arquivo"], "") if it["arquivo"] else ""
            if url:
                arq = f'<a href="{url}" target="_blank">abrir ↗</a>'
            elif it["arquivo"]:
                arq = f'<span class="arq">{it["arquivo"]}</span>'
            else:
                arq = "—"
            linhas += (f'<tr><td>{it["nome"]}</td>'
                       f'<td style="color:{c};white-space:nowrap">{ic} {txt}</td><td>{arq}</td></tr>')
        tabelas += (f'<div class="card"><h2>{e["papel"]}</h2>'
                    f'<table class="rel"><tr><th>Documento / Certidão</th><th>Situação</th><th>Arquivo na pasta</th></tr>'
                    f'{linhas}</table></div>')

    crit = ""
    for txt, val, obs in criterios:
        if val is True:
            resp = '<b style="color:#1a7d3c">SIM</b>'
        elif val is False:
            resp = '<b style="color:#c0392b">NÃO</b>'
        else:
            resp = f'<b style="color:#b8860b">Conferir</b>{(" — " + obs) if obs else ""}'
        crit += f'<tr><td>{txt}</td><td style="white-space:nowrap">{resp}</td></tr>'

    if d["recomendacoes"]:
        recs = "<p class='sub'>Recomendações:</p><ol>" + "".join(f"<li>{r}</li>" for r in d["recomendacoes"]) + "</ol>"
    elif d["risco"] == "MÉDIO":
        recs = f"<p>{d['concl_texto']}</p>"
    else:
        recs = "<p>Não há impedimentos — pode seguir com a contratação.</p>"

    corpo = f"""
    <div class="card cab">
      <h1>Relatório de Due Diligence{idtxt}</h1>
      <p class="titulo">{d['titulo']}</p>
      <span class="badge" style="background:{cor}">Risco: {d['risco']}</span>
      <p class="data">Gerado em {d['gerado_em']}</p>
    </div>
    <div class="card"><h2>Identificação</h2>{ident}</div>
    {tabelas}
    <div class="card"><h2>Critérios de risco</h2>
      <table class="rel"><tr><th>Critério</th><th>Resposta</th></tr>{crit}</table>
    </div>
    <div class="card"><h2>Conclusão</h2>
      <p>Risco: <b style="color:{cor}">{d['risco']}</b>.</p>{recs}
      <p>📄 <b>Parecer Jurídico:</b> arquivo <b>{ARQ_PARECER}</b> (na mesma pasta da DD).</p>
      {pasta_btn}
    </div>
    <div class="card" style="background:#eef6ff"><b>Para compartilhar:</b> envie o link da
      pasta da DD no Google Drive (botão acima) — ela contém este relatório, o parecer e todas as certidões.</div>
    <p class="rodape">Documento gerado automaticamente — revise antes de encaminhar ao setor de Franquias.</p>
    """
    return (f'<!doctype html><html lang="pt-br"><head><meta charset="utf-8">'
            f'<title>Relatório DD — {d["titulo"]}</title>'
            f"<style>{_RELATORIO_ESTILO}</style></head><body>{corpo}</body></html>")
